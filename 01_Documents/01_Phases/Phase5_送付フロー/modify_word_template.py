
import os
import shutil
from datetime import datetime
import docx
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# --- Configuration ---
base_dir = os.path.dirname(os.path.abspath(__file__))
templates_dir = os.path.abspath(os.path.join(base_dir, "../../02_Templates"))
target_template_path = os.path.join(templates_dir, "発注書テンプレート.docx")
backup_dir = os.path.join(base_dir, "Backups")
logo_path = os.path.join(base_dir, "TestOutput", "logo_c.png")

# --- Constants ---
COMPANY_INFO = [
    "千葉県千葉市中央区亥鼻1-8-15",
    "千葉大亥鼻イノベーションプラザ208",
    "セルジェンテック株式会社",
    "担当：　千賀　聡志",
    "",
    "TEL：FAX　043-441-4121"
]

GREETING_TEXT = (
    "平素よりお世話になっております。下記の通り発注させて頂きますのでご確認下さい。\n"
    "下記、発注内容で受注頂ける場合はお手数ですが発注請書に納品予定日、金額をご入力の上、３日以内にご返送ください。\n"
    "何卒、宜しくお願い申し上げます。"
)

ORDERER_TEXT = "発注依頼者: <<Orderer>>"

# --- Helpers ---

def create_backup(file_path):
    if not os.path.exists(backup_dir):
        os.makedirs(backup_dir)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M")
    filename = os.path.basename(file_path)
    name, ext = os.path.splitext(filename)
    backup_path = os.path.join(backup_dir, f"{name}_{timestamp}{ext}")
    shutil.copy2(file_path, backup_path)
    print(f"Backup created: {backup_path}")

def add_content_to_header_obj(header_obj, logo_path):
    # Check if we already added a table? (rudimentary check: count tables)
    # But checking content is hard. Assuming we modify clean template or idempotent-ish.
    
    table = header_obj.add_table(rows=1, cols=2, width=Mm(170))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    # Left Cell: Logo
    cell_left = table.cell(0, 0)
    cell_left.width = Mm(40)
    p_left = cell_left.paragraphs[0]
    run_left = p_left.add_run()
    if os.path.exists(logo_path):
        run_left.add_picture(logo_path, width=Mm(25))
    else:
        run_left.add_text("[LOGO MISSING]")
    
    # Right Cell: Company Info
    cell_right = table.cell(0, 1)
    cell_right.width = Mm(130)
    # Clear default paragraph
    cell_right._element.remove(cell_right.paragraphs[0]._element)
    
    for line in COMPANY_INFO:
        p = cell_right.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        
        # Safe font sizing
        run = p.runs[0] if p.runs else p.add_run()
        run.font.size = Pt(9)

    # Move table to the top of the header XML
    if len(header_obj._element) > 1:
        tbl = table._element
        header_obj._element.remove(tbl)
        header_obj._element.insert(0, tbl)

def find_sdt_by_tag_or_alias(parent_element, target_val):
    found = []
    for sdt in parent_element.iter(qn('w:sdt')):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is not None:
            tag = sdtPr.find(qn('w:tag'))
            alias = sdtPr.find(qn('w:alias'))
            
            tag_match = tag is not None and tag.get(qn('w:val')) == target_val
            alias_match = alias is not None and alias.get(qn('w:val')) == target_val
            
            if tag_match or alias_match:
                found.append(sdt)
    return found

# --- Main Logic ---

print(f"Processing template: {target_template_path}")
if not os.path.exists(target_template_path):
    print("Error: Template not found.")
    exit(1)

create_backup(target_template_path)
doc = docx.Document(target_template_path)

# 1. Header Modification
print("Modifying Headers...")
for i, section in enumerate(doc.sections):
    print(f"  Section {i+1}:")
    has_diff_first = section.different_first_page_header_footer
    has_odd_even = doc.settings.odd_and_even_pages_header_footer

    # Primary
    if not section.header.is_linked_to_previous or i == 0:
        print("    - Updating Standard Header")
        add_content_to_header_obj(section.header, logo_path)
    
    # First Page
    if has_diff_first:
        if not section.first_page_header.is_linked_to_previous or i == 0:
             print("    - Updating First Page Header")
             add_content_to_header_obj(section.first_page_header, logo_path)
    
    # Even Page
    if has_odd_even:
        if not section.even_page_header.is_linked_to_previous or i == 0:
            print("    - Updating Even Page Header")
            add_content_to_header_obj(section.even_page_header, logo_path)


# 2. Greeting Text
print("Inserting Greeting Text...")
greeting_p = doc.add_paragraph(GREETING_TEXT)
greeting_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
body = doc._element.body
p_element = greeting_p._element
body.remove(p_element)

# Insert after index 1 (Assuming Title/Spacer is 0) to avoid being top-most if Title exists
if len(body) > 0:
    body.insert(1, p_element)
else:
    body.insert(0, p_element)


# 3. Orderer Injection
print("Injecting Orderer field...")
target_tag = "DeliveryAddress"
found_sdts = find_sdt_by_tag_or_alias(doc._element.body, target_tag)

if not found_sdts:
    print("WARNING: DeliveryAddress not found. Appending to end.")
    doc.add_paragraph(ORDERER_TEXT)
else:
    sdt = found_sdts[0]
    parent = sdt.getparent()
    
    if parent.tag.endswith('p'): 
        p_parent = parent.getparent()
        index = list(p_parent).index(parent)
        
        temp_p = doc.add_paragraph(ORDERER_TEXT)
        temp_elm = temp_p._element
        doc._element.body.remove(temp_elm) 
        
        # Insert after target paragraph
        p_parent.insert(index + 1, temp_elm)
        print("Inserted after paragraph.")
        
    elif parent.tag.endswith('tc'): 
        # Inside Table Cell
        # Iterate up to find the w:tc element
        current = parent
        while current is not None and not current.tag.endswith('tc'):
            current = current.getparent()
            
        if current is not None:
            # Found TC
            # Append paragraph to TC
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = ORDERER_TEXT
            r.append(t)
            p.append(r)
            current.append(p)
            print("Inserted into table cell.")
        else:
            print("Could not find table cell parent. Fallback append.")
            doc.add_paragraph(ORDERER_TEXT)

doc.save(target_template_path)
print("Done.")
