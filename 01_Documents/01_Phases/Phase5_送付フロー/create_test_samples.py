
import os
import copy
import docx
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import openpyxl

# --- Configuration ---
base_dir = os.path.dirname(os.path.abspath(__file__))
templates_dir = os.path.abspath(os.path.join(base_dir, "../../02_Templates"))
word_template_path = os.path.join(templates_dir, "発注書テンプレート.docx")
excel_template_path = os.path.join(templates_dir, "請書テンプレート.xlsx")
output_dir = os.path.join(base_dir, "TestOutput")
word_output_path = os.path.join(output_dir, "発注書_検証用_MultiItem.docx")
excel_output_path = os.path.join(output_dir, "請書_検証用_MultiItem.xlsx")

# --- Dummy Data ---
DUMMY_DATA = {
    "Title": "検体解析業務委託",
    "OrderDate": datetime.now().strftime("%Y/%m/%d"),
    "VendorName": "テストサプライヤー株式会社",
    "DeliveryAddress": "千葉県千葉市中央区亥鼻1-8-15\n千葉大亥鼻イノベーションプラザ208",
    "RequestorName": "千賀 聡志",
}

ITEMS = [
    {
        "QuoteNumber": "Q-20260206-001",
        "ItemName": "DNA抽出キット",
        "Manufacturer": "メーカーA",
        "Quantity": 2,
        "UnitPrice": 50000,
    },
    {
        "QuoteNumber": "Q-20260206-002",
        "ItemName": "NGS解析サービス",
        "Manufacturer": "メーカーB",
        "Quantity": 1,
        "UnitPrice": 900000,
    },
    {
        "QuoteNumber": "",
        "ItemName": "サンプル輸送費",
        "Manufacturer": "メーカーC",
        "Quantity": 1,
        "UnitPrice": 15000,
    },
    {
        "QuoteNumber": "Q-20260206-004",
        "ItemName": "試薬A",
        "Manufacturer": "メーカーD",
        "Quantity": 5,
        "UnitPrice": 2000,
    },
    {
        "QuoteNumber": "Q-20260206-005",
        "ItemName": "試薬B",
        "Manufacturer": "メーカーE",
        "Quantity": 10,
        "UnitPrice": 1500,
    },
]

# --- Helper Functions (Word) ---

def set_sdt_text(sdt, text):
    content = sdt.find(qn('w:sdtContent'))
    if content is None: return
    # Remove existing
    for child in list(content):
        content.remove(child)
    # Add new content respecting inline vs block SDT
    parent = sdt.getparent()
    if parent is not None and parent.tag == qn('w:p'):
        # Inline SDT: content must be runs, not paragraphs
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = str(text)
        r.append(t)
        content.append(r)
    else:
        # Block SDT: content should be paragraphs
        p = OxmlElement('w:p')
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = str(text)
        r.append(t)
        p.append(r)
        content.append(p)

def populate_word_sdts(doc, data):
    for sdt in doc.element.body.iter(qn('w:sdt')):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is None: continue
        tag = sdtPr.find(qn('w:tag'))
        if tag is None: continue
        
        val = tag.get(qn('w:val'))
        if val in data:
            set_sdt_text(sdt, data[val])

def iter_sdt_tags(element):
    for sdt in element.iter(qn('w:sdt')):
        sdtPr = sdt.find(qn('w:sdtPr'))
        if sdtPr is None:
            continue
        tag = sdtPr.find(qn('w:tag'))
        if tag is None:
            continue
        val = tag.get(qn('w:val'))
        if val:
            yield val

def find_item_table(doc):
    for table in doc.tables:
        table_text = "\n".join(cell.text for row in table.rows for cell in row.cells)
        tags = set(iter_sdt_tags(table._element))
        if "品目" in table_text and {"ItemName", "Quantity"}.issubset(tags):
            return table
    return None

def find_item_template_row(table):
    for row in table.rows:
        tags = set(iter_sdt_tags(row._element))
        if {"QuoteNumber", "ItemName", "EstimatedAmount"}.issubset(tags):
            return row
    return None

# Helper to fill and FLATTEN SDT in a specific element
# Flattening means: remove w:sdt wrapper, keep content (w:p or w:r).
# This prevents ID duplication and corruption in cloned rows.
def flatten_row_sdts(row_element, data):
    # Iterate over all SDTs in the row
    # We must collect them first because we are modifying the tree
    sdts = list(row_element.iter(qn('w:sdt')))
    
    for sdt in sdts:
        sdtPr = sdt.find(qn('w:sdtPr'))
        tag_val = None
        if sdtPr is not None:
            tag = sdtPr.find(qn('w:tag'))
            if tag is not None:
                tag_val = tag.get(qn('w:val'))
        
        # Determine replacement text
        text_to_set = None
        if tag_val and tag_val in data:
            text_to_set = data[tag_val]
        
        # If we have data, we set it. If not, we keep original text (or empty?). 
        # But we MUST flatten to avoid ID issues.
        
        content = sdt.find(qn('w:sdtContent'))
        if content is None: continue
        
        # If we need to set text, we replace the content's text run
        if text_to_set is not None:
            # Simple approach: remove all children of content, add new run
            for child in list(content):
                content.remove(child)
            
            # Helper to create run
            # Check context: are we in a paragraph? or block level?
            # sdtContent can contain paragraphs (if block sdt) or runs (if inline sdt)
            parent = sdt.getparent()
            
            # If parent is w:p, sdt is inline -> content should contain w:r
            # If parent is w:tc, sdt is block -> content should contain w:p
            
            if parent.tag == qn('w:p'):
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = str(text_to_set)
                r.append(t)
                content.append(r)
            else:
                p = OxmlElement('w:p')
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = str(text_to_set)
                r.append(t)
                p.append(r)
                content.append(p)
                
        # FLATTEN: Move children of sdtContent to sdt's parent, then remove sdt
        parent = sdt.getparent()
        index = list(parent).index(sdt)
        
        for child in list(content):
            parent.insert(index, child)
            index += 1
            
        parent.remove(sdt)

def process_word_template(doc, data, items):
    # 1. Scalar Fields
    populate_word_sdts(doc, data)
    
    # 2. <<Orderer>> Replacement
    orderer_key = "<<Orderer>>"
    for p in doc.paragraphs:
        if orderer_key in p.text:
            p.text = p.text.replace(orderer_key, data.get("RequestorName", ""))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if orderer_key in p.text:
                        p.text = p.text.replace(orderer_key, data.get("RequestorName", ""))
    
    # 3. Tables (Items)
    table = find_item_table(doc)
    if table:
        print("  Found Item Table.")
        template_row = find_item_template_row(table)
        if template_row is None:
            print("  Error: Item template row not found.")
            return

        insert_index = list(table._element).index(template_row._element)
        table._element.remove(template_row._element)

        for offset, item in enumerate(items):
            # Calculate Data
            qty = item["Quantity"]
            price = item["UnitPrice"]
            amount = qty * price
            
            # Clone Row
            new_row_element = copy.deepcopy(template_row._element)
            
            row_data = {
                "QuoteNumber": str(item.get("QuoteNumber", "")),
                "ItemName": str(item["ItemName"]),
                "Manufacturer": str(item.get("Manufacturer", "")),
                "Quantity": str(qty),
                "EstimatedAmount": f"{amount:,}",
            }
            
            # CRITICAL FIX: Flatten SDTs in the cloned row
            flatten_row_sdts(new_row_element, row_data)
            
            table._element.insert(insert_index + offset, new_row_element)
    else:
        print("  WARNING: Item table not found in Word.")
                            


# --- Helper Functions (Excel) ---

def copy_row_style(ws, src_row_idx, dest_row_idx):
    src_row = ws[src_row_idx]
    dest_row = ws[dest_row_idx]
    
    for i, cell in enumerate(src_row):
        new_cell = dest_row[i]
        new_cell.font = copy.copy(cell.font)
        new_cell.border = copy.copy(cell.border)
        new_cell.fill = copy.copy(cell.fill)
        new_cell.number_format = cell.number_format
        new_cell.protection = copy.copy(cell.protection)
        new_cell.alignment = copy.copy(cell.alignment)

def find_excel_layout(ws):
    start_row = None
    subtotal_row = None
    for row in range(1, ws.max_row + 1):
        if ws.cell(row=row, column=2).value == "{{ItemName}}":
            start_row = row
        if ws.cell(row=row, column=5).value == "小計":
            subtotal_row = row
            break
    if start_row is None:
        start_row = 10
    if subtotal_row is None:
        subtotal_row = ws.max_row + 1
    data_end_row = subtotal_row - 1
    return start_row, data_end_row, subtotal_row

def process_excel_template(wb, data, items):
    ws = wb.active

    # Header fields
    ws["B3"] = data.get("VendorName", ws["B3"].value)
    ws["G5"] = data.get("OrderDate", ws["G5"].value)

    start_row, data_end_row, subtotal_row = find_excel_layout(ws)
    available_rows = data_end_row - start_row + 1

    # Expand item area only when template rows are insufficient
    if len(items) > available_rows:
        extra = len(items) - available_rows
        insert_at = data_end_row + 1
        for _ in range(extra):
            ws.insert_rows(insert_at)
            copy_row_style(ws, start_row, insert_at)
            data_end_row += 1
            subtotal_row += 1

    for row in range(start_row, data_end_row + 1):
        idx = row - start_row
        ws.cell(row=row, column=1).value = idx + 1
        ws.cell(row=row, column=5).value = f'=IF(D{row}<>"",F{row}/D{row},"")'

        if idx < len(items):
            item = items[idx]
            qty = item["Quantity"]
            amount = qty * item["UnitPrice"]
            ws.cell(row=row, column=2).value = item["ItemName"]
            ws.cell(row=row, column=3).value = item.get("Manufacturer", "")
            ws.cell(row=row, column=4).value = qty
            ws.cell(row=row, column=6).value = amount
            ws.cell(row=row, column=7).value = None
            ws.cell(row=row, column=8).value = item.get("QuoteNumber", "")
        else:
            ws.cell(row=row, column=2).value = None
            ws.cell(row=row, column=3).value = None
            ws.cell(row=row, column=4).value = None
            ws.cell(row=row, column=6).value = None
            ws.cell(row=row, column=7).value = None
            ws.cell(row=row, column=8).value = None

    # Keep summary formulas aligned with the current data range
    ws.cell(row=subtotal_row, column=6).value = f"=SUM(F{start_row}:F{data_end_row})"
    ws.cell(row=subtotal_row + 1, column=6).value = f"=F{subtotal_row}*0.1"
    ws.cell(row=subtotal_row + 2, column=6).value = f"=F{subtotal_row}+F{subtotal_row + 1}"

# --- Main ---

def main():
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    blank_quote_count = sum(1 for item in ITEMS if not str(item.get("QuoteNumber", "")).strip())
    print(f"QuoteNumber blank count: {blank_quote_count}")

    # --- Process Word ---
    print(f"Processing Word Template: {word_template_path}")
    if os.path.exists(word_template_path):
        doc = docx.Document(word_template_path)
        process_word_template(doc, DUMMY_DATA, ITEMS)
        doc.save(word_output_path)
        print(f"  Saved Word to: {word_output_path}")
    else:
        print("  Word template not found.")

    # --- Process Excel ---
    print(f"Processing Excel Template: {excel_template_path}")
    if os.path.exists(excel_template_path):
        wb = openpyxl.load_workbook(excel_template_path)
        process_excel_template(wb, DUMMY_DATA, ITEMS)
        wb.save(excel_output_path)
        print(f"  Saved Excel to: {excel_output_path}")
    else:
        print("  Excel template not found.")

if __name__ == "__main__":
    main()
