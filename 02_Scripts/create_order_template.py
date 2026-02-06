"""
発注書Wordテンプレート作成スクリプト
Power Automate「Populate a Microsoft Word template」アクション用のContent Control付きテンプレートを生成

Content Control Names（SP内部名に対応）:
- Title: 件名
- ItemName: 品目
- Manufacturer: メーカー
- Quantity: 数量
- EstimatedAmount: 見積額
- DeliveryAddress: 納品先
- QuoteNumber: 見積書番号
- VendorName: 業者名
- OrderDate: 発注日
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import os

def add_content_control(run, tag, alias=None):
    """
    テキストにContent Control (Structured Document Tag)を追加
    Power Automateが認識する形式で挿入
    """
    # SDT要素を作成
    sdt = OxmlElement('w:sdt')
    
    # SDT Properties
    sdtPr = OxmlElement('w:sdtPr')
    
    # Alias (表示名)
    if alias:
        sdtAlias = OxmlElement('w:alias')
        sdtAlias.set(qn('w:val'), alias)
        sdtPr.append(sdtAlias)
    
    # Tag (タグ名 - SP内部名に対応)
    sdtTag = OxmlElement('w:tag')
    sdtTag.set(qn('w:val'), tag)
    sdtPr.append(sdtTag)
    
    # ID
    sdtId = OxmlElement('w:id')
    sdtId.set(qn('w:val'), str(abs(hash(tag)) % (10 ** 8)))
    sdtPr.append(sdtId)
    
    # Lock（編集ロック：コンテンツのみ変更許可）
    lock = OxmlElement('w:lock')
    lock.set(qn('w:val'), 'sdtContentLocked')
    sdtPr.append(lock)
    
    # Placeholder（プレースホルダーテキスト）
    showingPlaceholder = OxmlElement('w:showingPlcHdr')
    sdtPr.append(showingPlaceholder)
    
    # Plain Text Content Control (必須)
    text = OxmlElement('w:text')
    sdtPr.append(text)
    
    sdt.append(sdtPr)
    
    # SDT Content（プレースホルダーテキスト）
    sdtContent = OxmlElement('w:sdtContent')
    
    # プレースホルダー用のrun要素
    placeholderRun = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    # プレースホルダースタイル
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'PlaceholderText')
    rPr.append(rStyle)
    placeholderRun.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = f"《{tag}》"
    placeholderRun.append(t)
    sdtContent.append(placeholderRun)
    
    sdt.append(sdtContent)
    
    # 親要素に挿入
    run._r.getparent().replace(run._r, sdt)
    
    return sdt

def create_order_template():
    """発注書Wordテンプレートを作成"""
    doc = Document()
    
    # 文書のプロパティ設定
    sections = doc.sections
    for section in sections:
        section.page_height = Cm(29.7)  # A4
        section.page_width = Cm(21)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    
    # タイトル
    title = doc.add_heading('発 注 書', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 発注日
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run('発注日：')
    add_content_control(date_para.add_run('《OrderDate》'), 'OrderDate', '発注日')
    
    doc.add_paragraph()  # スペース
    
    # 宛先（業者名）
    vendor_para = doc.add_paragraph()
    vendor_run = vendor_para.add_run('')
    add_content_control(vendor_para.add_run('《VendorName》'), 'VendorName', '業者名')
    vendor_para.add_run('　御中')
    
    doc.add_paragraph()  # スペース
    
    # 発注者情報
    sender_para = doc.add_paragraph()
    sender_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sender_para.add_run('セルジェンテック株式会社')
    
    doc.add_paragraph()  # スペース
    
    # 注文内容表
    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # ヘッダー行
    header_cells = table.rows[0].cells
    headers = ['件名', '品目', 'メーカー', '数量', '見積額']
    for i, text in enumerate(headers):
        header_cells[i].text = text
        header_cells[i].paragraphs[0].runs[0].bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # データ行（Content Control）
    data_row = table.rows[1].cells
    
    # Title
    p = data_row[0].paragraphs[0]
    add_content_control(p.add_run('《Title》'), 'Title', '件名')
    
    # ItemName
    p = data_row[1].paragraphs[0]
    add_content_control(p.add_run('《ItemName》'), 'ItemName', '品目')
    
    # Manufacturer
    p = data_row[2].paragraphs[0]
    add_content_control(p.add_run('《Manufacturer》'), 'Manufacturer', 'メーカー')
    
    # Quantity
    p = data_row[3].paragraphs[0]
    add_content_control(p.add_run('《Quantity》'), 'Quantity', '数量')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # EstimatedAmount
    p = data_row[4].paragraphs[0]
    add_content_control(p.add_run('《EstimatedAmount》'), 'EstimatedAmount', '見積額')
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 見積書番号行
    quote_row = table.rows[2].cells
    quote_row[0].text = '見積書番号'
    quote_row[0].paragraphs[0].runs[0].bold = True
    
    # QuoteNumber（セル結合）
    merged_cell = quote_row[1].merge(quote_row[4])
    p = merged_cell.paragraphs[0]
    p.clear()
    add_content_control(p.add_run('《QuoteNumber》'), 'QuoteNumber', '見積書番号')
    
    doc.add_paragraph()  # スペース
    
    # 納品先
    delivery_para = doc.add_paragraph()
    delivery_para.add_run('納品先：　')
    add_content_control(delivery_para.add_run('《DeliveryAddress》'), 'DeliveryAddress', '納品先')
    
    doc.add_paragraph()  # スペース
    
    # 備考
    doc.add_paragraph('【備考】')
    doc.add_paragraph('・納期については別途ご相談ください。')
    doc.add_paragraph('・本発注書と請書（返送用）をお送りいたします。')
    doc.add_paragraph('・ご不明点がございましたらご連絡ください。')
    
    return doc

if __name__ == '__main__':
    output_path = r'c:\Users\千賀聡志\OneDrive - セルジェンテック株式会社\sharepointリスト化\Documents\02_Templates\発注書テンプレート.docx'
    
    doc = create_order_template()
    doc.save(output_path)
    
    print(f'✅ 発注書テンプレートを作成しました: {output_path}')
    print()
    print('Content Control一覧:')
    print('  - Title: 件名')
    print('  - ItemName: 品目')
    print('  - Manufacturer: メーカー')
    print('  - Quantity: 数量')
    print('  - EstimatedAmount: 見積額')
    print('  - DeliveryAddress: 納品先')
    print('  - QuoteNumber: 見積書番号')
    print('  - VendorName: 業者名')
    print('  - OrderDate: 発注日')
